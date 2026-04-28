"""
Экспорт результата транскрибации в разные форматы.

Текстовые форматы (txt/srt/vtt/json) возвращаются как строки.
DOCX — бинарный, возвращается как bytes.

DOCX-варианты:
  - plain       — просто текст, без таймкодов
  - timecoded   — [00:01:15] Текст сегмента
  - full        — [00:01:15] SPEAKER_00: Текст сегмента
"""
from __future__ import annotations

import io
import json
from enum import Enum

from docx import Document
from docx.shared import Pt

from app.schemas.transcription import Segment, TranscriptionResult


class ExportFormat(str, Enum):
    """Текстовые форматы."""
    TXT = "txt"
    SRT = "srt"
    VTT = "vtt"
    JSON = "json"


class DocxStyle(str, Enum):
    """Стили представления DOCX."""
    PLAIN = "plain"          # только текст
    TIMECODED = "timecoded"  # [время] текст
    FULL = "full"            # [время] Спикер: текст


# ---------- Текстовые форматы (возвращают str) ----------

def export_text(result: TranscriptionResult, fmt: ExportFormat) -> tuple[str, str]:
    """Возвращает (содержимое, MIME-тип) для текстовых форматов."""
    if fmt is ExportFormat.TXT:
        return _to_txt(result), "text/plain; charset=utf-8"
    if fmt is ExportFormat.SRT:
        return _to_srt(result.segments), "application/x-subrip; charset=utf-8"
    if fmt is ExportFormat.VTT:
        return _to_vtt(result.segments), "text/vtt; charset=utf-8"
    if fmt is ExportFormat.JSON:
        return _to_json(result), "application/json; charset=utf-8"
    raise ValueError(f"Unknown format: {fmt}")


def _to_txt(result: TranscriptionResult) -> str:
    lines = [seg.text.strip() for seg in result.segments if seg.text.strip()]
    return "\n\n".join(lines) + "\n"


def _to_srt(segments: list[Segment]) -> str:
    blocks: list[str] = []
    for i, seg in enumerate(segments, start=1):
        blocks.append(
            f"{i}\n"
            f"{_format_time_srt(seg.start)} --> {_format_time_srt(seg.end)}\n"
            f"{seg.text.strip()}\n"
        )
    return "\n".join(blocks)


def _format_time_srt(seconds: float) -> str:
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds - int(seconds)) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"


def _to_vtt(segments: list[Segment]) -> str:
    blocks = ["WEBVTT\n"]
    for seg in segments:
        blocks.append(
            f"{_format_time_vtt(seg.start)} --> {_format_time_vtt(seg.end)}\n"
            f"{seg.text.strip()}\n"
        )
    return "\n".join(blocks)


def _format_time_vtt(seconds: float) -> str:
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds - int(seconds)) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"


def _to_json(result: TranscriptionResult) -> str:
    return json.dumps(result.model_dump(mode="json"), ensure_ascii=False, indent=2)


# ---------- DOCX (возвращает bytes) ----------

def export_docx(
    result: TranscriptionResult,
    style: DocxStyle,
    title: str | None = None,
) -> bytes:
    """
    Генерирует Word-документ в памяти и возвращает bytes.

    Args:
        result: что экспортируем
        style: один из трёх стилей представления
        title: заголовок в начале документа (обычно имя файла).

    Returns:
        Содержимое .docx в виде bytes, готовое к записи или отправке.
    """
    doc = Document()

    # Базовый стиль — 11pt Calibri, стандартно для современных Word-документов
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    if title:
        doc.add_heading(title, level=1)

    # Подзаголовок с метаданными — общие для всех стилей
    meta = doc.add_paragraph()
    meta.add_run(
        f"Язык: {result.language}    "
        f"Длительность: {_human_duration(result.duration_sec)}    "
        f"Сегментов: {len(result.segments)}"
    ).italic = True

    doc.add_paragraph()  # пустая строка-разделитель

    if style is DocxStyle.PLAIN:
        _docx_plain(doc, result)
    elif style is DocxStyle.TIMECODED:
        _docx_timecoded(doc, result)
    elif style is DocxStyle.FULL:
        _docx_full(doc, result)
    else:
        raise ValueError(f"Unknown DOCX style: {style}")

    # Сохраняем в буфер и возвращаем bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _docx_plain(doc: "Document", result: TranscriptionResult) -> None:
    """Чистый текст — один сегмент = один абзац."""
    for seg in result.segments:
        text = seg.text.strip()
        if text:
            doc.add_paragraph(text)


def _docx_timecoded(doc: "Document", result: TranscriptionResult) -> None:
    """[00:01:15] Текст сегмента"""
    for seg in result.segments:
        text = seg.text.strip()
        if not text:
            continue
        p = doc.add_paragraph()
        # Таймкод жирным, чтобы легко выцепить глазами
        p.add_run(f"[{_format_time_docx(seg.start)}] ").bold = True
        p.add_run(text)


def _docx_full(doc: "Document", result: TranscriptionResult) -> None:
    """[00:01:15] SPEAKER_00: Текст"""
    for seg in result.segments:
        text = seg.text.strip()
        if not text:
            continue
        p = doc.add_paragraph()
        p.add_run(f"[{_format_time_docx(seg.start)}] ").bold = True
        speaker = seg.speaker or "SPEAKER_?"
        p.add_run(f"{speaker}: ").bold = True
        p.add_run(text)


def _format_time_docx(seconds: float) -> str:
    """Человеко-читаемый таймкод: HH:MM:SS (без миллисекунд — в DOCX не нужны)."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    if hours:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"


def _human_duration(seconds: float) -> str:
    """«14 мин 22 сек» и подобное, для шапки документа."""
    total = int(seconds)
    h, m, s = total // 3600, (total % 3600) // 60, total % 60
    if h:
        return f"{h} ч {m} мин {s} сек"
    if m:
        return f"{m} мин {s} сек"
    return f"{s} сек"
